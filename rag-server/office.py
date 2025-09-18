import os
from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())

import logging
from mcp.server.fastmcp import FastMCP

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

from docx import Document as WordDocument
import pandas as pd


logging.basicConfig(level=logging.INFO)

mcp = FastMCP("Office-RAG")

OFFICE_DIR = os.getenv("OFFICE_DIR")

def load_office_documents(folder_path:str) -> list[Document]:
    docs = []

    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)

        if filename.endswith(".docx"):
            word = WordDocument(path)
            full_text = "\n".join([p.text for p in word.paragraphs if p.text.strip()])
            docs.append(Document(page_content=full_text, metadata={"source": filename}))
        elif filename.endswith(".xlsx"):
            try:
                excel = pd.read_excel(path, sheet_name=None)
                for sheet_name, df in excel.items():
                    text = df.astype(str).to_string(index=False)
                    docs.append(Document(page_content=text, metadata={"source": f"{filename} - {sheet_name}"}))
            except Exception as e:
                logging.error(f"Error loading Excel file {filename}: {e}")
    return docs

raw_docs = load_office_documents(OFFICE_DIR)
print(f"Loaded {len(raw_docs)} documents")

splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
docs = splitter.split_documents(raw_docs)
print(f"Split into {len(docs)} documents")

embeddings = OpenAIEmbeddings()
llm = ChatOpenAI(model="gpt-4.1-nano", temperature=0)

vectorestore = Chroma.from_documents(
    documents=docs,
    embedding=embeddings,
)

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorestore.as_retriever(),
)

@mcp.tool()
def ask_office(query: str) -> str:
    """폴더 내 Word/Excel 문서를 기반으로 질문에 답변합니다."""
    logging.info(f"Received query: {query}")
    response = qa_chain.run(query)
    return response

if __name__ == "__main__":
    mcp.run(transport="stdio")